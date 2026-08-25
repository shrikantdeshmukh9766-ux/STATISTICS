"""
Stream-lite — Baseline Table Builder
=====================================
A Streamlit app that turns an uploaded master chart (Excel/CSV) into a
publication-ready "Table 1": descriptive statistics per variable, with
automatic parametric/non-parametric test selection.

Run with:
    pip install streamlit pandas numpy scipy openpyxl
    streamlit run streamlite_app.py

How test selection works
-------------------------
Numeric variables, 2 groups   -> Welch's t-test (parametric) or
                                  Mann-Whitney U / Wilcoxon rank-sum (non-parametric)
Numeric variables, 3+ groups  -> One-way ANOVA or Kruskal-Wallis H
Categorical variables         -> Chi-square test of independence, automatically
                                  switched to Fisher's exact test for 2x2 tables
                                  when any expected cell count is below 5

Parametric vs non-parametric is decided automatically via the
D'Agostino-Pearson omnibus normality test (scipy.stats.normaltest),
unless the user forces one or the other.
"""

import io
import numpy as np
import pandas as pd
import streamlit as st
from scipy import stats
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

st.set_page_config(page_title="Stream-lite · Baseline Table Builder", layout="wide")

# --------------------------------------------------------------------------
# Statistical helpers
# --------------------------------------------------------------------------

def detect_type(series: pd.Series):
    """Guess whether a column is numerical or categorical."""
    non_missing = series.dropna()
    non_missing = non_missing[non_missing.astype(str).str.strip() != ""]
    n = len(non_missing)
    if n == 0:
        return "categorical", 0, 0
    numeric_coerced = pd.to_numeric(non_missing, errors="coerce")
    numeric_ratio = numeric_coerced.notna().mean()
    unique_n = non_missing.astype(str).str.strip().nunique()
    if numeric_ratio >= 0.9 and unique_n > 10:
        return "numerical", n, unique_n
    return "categorical", n, unique_n


def is_normal(arr, alpha):
    """D'Agostino-Pearson omnibus normality test. Needs n>=8; smaller
    samples are treated as non-normal (safer default)."""
    arr = np.asarray(arr, dtype=float)
    arr = arr[~np.isnan(arr)]
    if len(arr) < 8:
        return False
    if np.all(arr == arr[0]):
        return True
    try:
        _, p = stats.normaltest(arr)
        return p > alpha
    except Exception:
        return False


def welch_t_test(a, b):
    a, b = np.asarray(a, float), np.asarray(b, float)
    res = stats.ttest_ind(a, b, equal_var=False)
    v1, v2, n1, n2 = a.var(ddof=1), b.var(ddof=1), len(a), len(b)
    df = (v1 / n1 + v2 / n2) ** 2 / ((v1 / n1) ** 2 / (n1 - 1) + (v2 / n2) ** 2 / (n2 - 1))
    return {"name": "Independent t-test (Welch)",
            "stat": f"t={res.statistic:.2f}, df={df:.1f}",
            "p": float(res.pvalue)}


def mann_whitney_test(a, b):
    res = stats.mannwhitneyu(a, b, alternative="two-sided")
    return {"name": "Mann-Whitney U (Wilcoxon rank-sum)",
            "stat": f"U={res.statistic:.1f}",
            "p": float(res.pvalue)}


def anova_test(groups):
    res = stats.f_oneway(*groups)
    k = len(groups)
    N = sum(len(g) for g in groups)
    return {"name": "One-way ANOVA",
            "stat": f"F={res.statistic:.2f}, df={k-1},{N-k}",
            "p": float(res.pvalue)}


def kruskal_test(groups):
    res = stats.kruskal(*groups)
    return {"name": "Kruskal-Wallis H",
            "stat": f"H={res.statistic:.2f}, df={len(groups)-1}",
            "p": float(res.pvalue)}


def chi_or_fisher_test(table):
    """RxC contingency table -> chi-square, auto-falling back to Fisher's
    exact test for 2x2 tables with low expected counts."""
    table = np.array(table)
    chi2, p, dof, expected = stats.chi2_contingency(table)
    min_e = float(expected.min())
    if table.shape == (2, 2) and min_e < 5:
        _, p_fisher = stats.fisher_exact(table)
        return {"name": "Fisher's exact test", "stat": "—", "p": float(p_fisher), "min_expected": min_e}
    return {"name": "Chi-square test", "stat": f"\u03C7\u00B2={chi2:.2f}, df={dof}",
            "p": float(p), "min_expected": min_e}


def fmt_p(p):
    if p is None or (isinstance(p, float) and np.isnan(p)):
        return "—"
    return "<0.001" if p < 0.001 else f"{p:.3f}"


def fmt_num(x, d=1):
    return f"{x:.{d}f}"


# --------------------------------------------------------------------------
# Table 1 builder
# --------------------------------------------------------------------------

def format_numeric_cell(values, display_mode, use_param):
    """Format one group's numeric summary according to the chosen display mode.
    display_mode: 'auto' | 'mean_sd' | 'median_iqr' | 'both'"""
    if len(values) == 0:
        return "—"
    show_mean = display_mode == "mean_sd" or display_mode == "both" or (display_mode == "auto" and use_param)
    show_median = display_mode == "median_iqr" or display_mode == "both" or (display_mode == "auto" and not use_param)
    parts = []
    if show_mean:
        parts.append(f"{fmt_num(np.mean(values))} \u00B1 {fmt_num(np.std(values, ddof=1))}")
    if show_median:
        q1, med, q3 = np.percentile(values, [25, 50, 75])
        parts.append(f"{fmt_num(med)} ({fmt_num(q1)}\u2013{fmt_num(q3)})")
    return "; ".join(parts)


def numeric_label(col, display_mode, use_param):
    if display_mode == "mean_sd":
        return f"{col}, mean \u00B1 SD"
    if display_mode == "median_iqr":
        return f"{col}, median (IQR)"
    if display_mode == "both":
        return f"{col}, mean \u00B1 SD; median (IQR)"
    return f"{col}, mean \u00B1 SD" if use_param else f"{col}, median (IQR)"


def build_table1(df, var_meta, group_col, test_mode, alpha, display_mode="auto"):
    """Returns (display_rows, csv_rows, footnote_flags).
    display_rows: list of dicts describing each printed row for st rendering.
    csv_rows: list of lists for CSV / clipboard export.
    display_mode controls how numeric summaries are shown: 'auto' (mean±SD if
    normal else median (IQR)), 'mean_sd', 'median_iqr', or 'both'. This is
    independent of test_mode, which controls whether t-test/ANOVA or
    Wilcoxon/Kruskal-Wallis is used for the comparison.
    """
    group_levels = []
    if group_col:
        group_levels = sorted(df[group_col].dropna().astype(str).str.strip().unique().tolist())

    header = ["Variable"]
    if group_col:
        for lv in group_levels:
            n = (df[group_col].astype(str).str.strip() == lv).sum()
            header.append(f"{lv} (n={n})")
        header += ["Test", "Statistic", "p"]
    else:
        header.append(f"Overall (n={len(df)})")

    csv_rows = [header]
    display_rows = []
    flags = set()

    variables = [c for c in df.columns if var_meta[c]["use"] and c != group_col]

    for col in variables:
        vtype = var_meta[col]["type"]

        if vtype == "numerical":
            numeric_series = pd.to_numeric(df[col], errors="coerce")

            if group_col:
                group_arrays = []
                for lv in group_levels:
                    mask = (df[group_col].astype(str).str.strip() == lv) & numeric_series.notna()
                    group_arrays.append(numeric_series[mask].values)
            else:
                group_arrays = None

            all_vals = numeric_series.dropna().values

            # use_param governs which TEST is used (t-test/ANOVA vs Wilcoxon/KW);
            # it is decided by test_mode regardless of the display_mode setting.
            if test_mode == "parametric":
                use_param = True
            elif test_mode == "nonparametric":
                use_param = False
            else:
                if group_col:
                    use_param = all(is_normal(g, alpha) for g in group_arrays if len(g) > 0)
                else:
                    use_param = is_normal(all_vals, alpha)

            label = numeric_label(col, display_mode, use_param)

            cells = []
            if group_col:
                for g in group_arrays:
                    cells.append(format_numeric_cell(g, display_mode, use_param))
            else:
                cells.append(format_numeric_cell(all_vals, display_mode, use_param))

            test_result = None
            if group_col and len(group_levels) >= 2:
                nonempty = [g for g in group_arrays if len(g) > 1]  # need >=2 points per group for variance
                if len(nonempty) >= 2:
                    try:
                        if len(group_levels) == 2:
                            test_result = welch_t_test(*nonempty) if use_param else mann_whitney_test(*nonempty)
                        else:
                            test_result = anova_test(nonempty) if use_param else kruskal_test(nonempty)
                    except Exception:
                        test_result = None
                        flags.add("skipped")

            display_rows.append({"kind": "var", "label": label, "cells": cells, "test": test_result})
            csv_rows.append([label, *cells,
                              test_result["name"] if test_result else "",
                              test_result["stat"] if test_result else "",
                              fmt_p(test_result["p"]) if test_result else ""])

        else:  # categorical
            series = df[col].astype(str).str.strip()
            series = series.where(df[col].notna() & (series != ""), other=np.nan)
            levels = sorted(series.dropna().unique().tolist())

            contingency = None
            test_result = None
            if group_col and len(group_levels) >= 2 and len(levels) >= 2:
                group_series = df[group_col].astype(str).str.strip()
                contingency = [[int(((series == lv) & (group_series == glv)).sum()) for glv in group_levels]
                               for lv in levels]
                # Skip the test if any row/column is entirely zero — chi2_contingency
                # (and Fisher's exact) require every row and column to have at least
                # one observation, otherwise the table is degenerate.
                arr = np.array(contingency)
                if arr.size > 0 and arr.shape[0] >= 2 and arr.shape[1] >= 2 \
                        and (arr.sum(axis=0) > 0).all() and (arr.sum(axis=1) > 0).all():
                    try:
                        res = chi_or_fisher_test(contingency)
                        test_result = res
                        if res["name"].startswith("Fisher"):
                            flags.add("fisher")
                        else:
                            flags.add("chi2")
                            if res["min_expected"] < 5:
                                flags.add("lowE")
                    except Exception:
                        test_result = None
                        flags.add("skipped")
            elif group_col:
                # Build a zero contingency table skeleton for the n(%) display below,
                # even though no test is run (fewer than 2 non-empty levels/groups).
                group_series = df[group_col].astype(str).str.strip()
                contingency = [[int(((series == lv) & (group_series == glv)).sum()) for glv in group_levels]
                               for lv in levels]

            display_rows.append({"kind": "varheader", "label": f"{col}, n (%)"})
            csv_rows.append([f"{col}, n (%)"])

            for i, lv in enumerate(levels):
                cells = []
                if group_col:
                    group_series = df[group_col].astype(str).str.strip()
                    for gi, glv in enumerate(group_levels):
                        n = contingency[i][gi]
                        gtotal = (group_series == glv).sum()
                        pct = 100 * n / gtotal if gtotal else 0.0
                        cells.append(f"{n} ({pct:.1f}%)")
                else:
                    n = int((series == lv).sum())
                    pct = 100 * n / len(df) if len(df) else 0.0
                    cells.append(f"{n} ({pct:.1f}%)")

                is_last = i == len(levels) - 1
                display_rows.append({"kind": "level", "label": lv, "cells": cells,
                                      "test": test_result if is_last else None})
                csv_rows.append([f"  {lv}", *cells,
                                  test_result["name"] if is_last and test_result else "",
                                  test_result["stat"] if is_last and test_result else "",
                                  fmt_p(test_result["p"]) if is_last and test_result else ""])

    return header, display_rows, csv_rows, flags


def render_table_markdown(header, display_rows, group_col, alpha):
    """Render the Table 1 as an HTML table with a three-line (journal-style) look."""
    css = """
    <style>
    table.pub { width:100%; border-collapse:collapse; font-family: Georgia, serif; font-size: 14px; }
    table.pub thead th { border-top:2px solid #1E2A32; border-bottom:1px solid #1E2A32;
                          padding:8px 10px; text-align:left; font-family: -apple-system, sans-serif; }
    table.pub tbody td { padding:5px 10px; }
    table.pub tbody tr.var td { font-weight:700; padding-top:10px; }
    table.pub tbody tr.level td.name { padding-left:20px; color:#555; font-weight:400; }
    table.pub td.stat { font-family: ui-monospace, Consolas, monospace; text-align:center; font-size:12.5px;}
    table.pub td.sig { font-weight:700; }
    table.pub tbody tr.lastrow td { border-bottom:2px solid #1E2A32; padding-bottom:10px; }
    </style>
    """
    html = css + '<table class="pub"><thead><tr>'
    for h in header:
        html += f"<th>{h}</th>"
    html += "</tr></thead><tbody>"

    for idx, row in enumerate(display_rows):
        is_last_of_block = (idx == len(display_rows) - 1) or \
                            (display_rows[idx + 1]["kind"] in ("var", "varheader"))
        cls = "var" if row["kind"] in ("var", "varheader") else "level"
        cls += " lastrow" if is_last_of_block else ""
        html += f'<tr class="{cls}">'

        if row["kind"] == "varheader":
            html += f'<td>{row["label"]}</td>'
            colspan = len(header) - 1
            html += f'<td colspan="{colspan}"></td>'
        else:
            name_cls = "name" if row["kind"] == "level" else ""
            html += f'<td class="{name_cls}">{row["label"]}</td>'
            for c in row["cells"]:
                html += f'<td class="stat">{c}</td>'
            if group_col:
                t = row.get("test")
                if t:
                    sig = "sig" if t["p"] < alpha else ""
                    html += f'<td>{t["name"]}</td><td class="stat">{t["stat"]}</td><td class="stat {sig}">{fmt_p(t["p"])}</td>'
                else:
                    html += "<td>—</td><td>—</td><td>—</td>"
        html += "</tr>"
    html += "</tbody></table>"
    return html


def _set_cell_border(cell, **kwargs):
    """Add borders to a single table cell. kwargs like
    top={'sz':12,'val':'single','color':'000000'}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            spec = kwargs[edge]
            tag = f'w:{edge}'
            el = tcBorders.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                tcBorders.append(el)
            el.set(qn('w:val'), spec.get('val', 'single'))
            el.set(qn('w:sz'), str(spec.get('sz', 8)))
            el.set(qn('w:color'), spec.get('color', '000000'))


def descriptive_footnote(display_mode, alpha):
    if display_mode == "mean_sd":
        return "Continuous variables reported as mean \u00B1 SD; categorical variables reported as n (%)."
    if display_mode == "median_iqr":
        return "Continuous variables reported as median (IQR); categorical variables reported as n (%)."
    if display_mode == "both":
        return "Continuous variables reported as mean \u00B1 SD and median (IQR); categorical variables reported as n (%)."
    return (f"Continuous variables reported as mean \u00B1 SD (assessed as normal via D'Agostino-Pearson "
            f"test, \u03B1={alpha}) or median (IQR) otherwise; categorical variables reported as n (%).")


def build_docx(header, display_rows, group_col, alpha, flags, display_mode="auto",
                title="Table 1. Baseline characteristics"):
    """Build a Word document with a three-line (journal-style) table."""
    doc = Document()
    h = doc.add_heading(title, level=2)
    for r in h.runs:
        r.font.size = Pt(13)

    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(header):
        hdr_cells[i].text = str(htext)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        _set_cell_border(hdr_cells[i], top={'sz': 12, 'val': 'single'},
                          bottom={'sz': 8, 'val': 'single'})

    n_rows = len(display_rows)
    for idx, row in enumerate(display_rows):
        cells = table.add_row().cells
        is_last = idx == n_rows - 1
        next_is_new_block = is_last or (display_rows[idx + 1]["kind"] in ("var", "varheader"))

        if row["kind"] == "varheader":
            cells[0].text = row["label"]
            for run in cells[0].paragraphs[0].runs:
                run.bold = True
            for c in cells[1:]:
                c.text = ""
        else:
            label = ("    " + row["label"]) if row["kind"] == "level" else row["label"]
            cells[0].text = label
            for run in cells[0].paragraphs[0].runs:
                run.bold = (row["kind"] == "var")
            ci = 1
            for val in row["cells"]:
                cells[ci].text = str(val)
                for p in cells[ci].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ci += 1
            if group_col:
                t = row.get("test")
                if t:
                    cells[ci].text = t["name"]; ci += 1
                    cells[ci].text = t["stat"]; ci += 1
                    cells[ci].text = fmt_p(t["p"])
                    for p in cells[ci].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if t["p"] < alpha:
                        for run in cells[ci].paragraphs[0].runs:
                            run.bold = True
                else:
                    cells[ci].text = ""; ci += 1
                    cells[ci].text = ""; ci += 1
                    cells[ci].text = ""

        if next_is_new_block:
            for c in cells:
                _set_cell_border(c, bottom={'sz': 8, 'val': 'single'})

    for c in table.rows[-1].cells:
        _set_cell_border(c, bottom={'sz': 12, 'val': 'single'})

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9.5)

    footnotes = [descriptive_footnote(display_mode, alpha)]
    if "chi2" in flags:
        footnotes.append("Chi-square test of independence used for categorical comparisons with adequate expected cell counts.")
    if "fisher" in flags:
        footnotes.append("Fisher's exact test used in place of chi-square when a 2\u00D72 table had an expected cell count below 5.")
    if "lowE" in flags:
        footnotes.append("Caution: one or more categorical comparisons above have expected cell counts below 5; chi-square approximation may be unreliable.")
    if "skipped" in flags:
        footnotes.append("A statistical test could not be computed for one or more variables and was left blank.")
    footnotes.append(f"Bold p-values indicate statistical significance at \u03B1={alpha}.")

    doc.add_paragraph()  # spacer
    for f in footnotes:
        p = doc.add_paragraph(f)
        for run in p.runs:
            run.font.size = Pt(8.5)
            run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# --------------------------------------------------------------------------
# Streamlit UI
# --------------------------------------------------------------------------

st.title("Stream-lite · Baseline Table Builder")
st.caption(
    "Upload a master chart. Assign variable types, choose a grouping variable, and "
    "Stream-lite auto-selects the correct test — t-test or Wilcoxon for numeric variables, "
    "chi-square or Fisher's exact for categorical — and lays out a three-line publication table."
)

st.markdown("### 1. Upload master chart")
uploaded = st.file_uploader("Excel (.xlsx/.xls) or CSV. First row must be column headers.",
                             type=["xlsx", "xls", "csv"])

if uploaded is not None:
    try:
        if uploaded.name.lower().endswith(".csv"):
            df = pd.read_csv(uploaded)
        else:
            df = pd.read_excel(uploaded)
    except Exception as e:
        st.error(f"Could not read that file: {e}")
        st.stop()

    st.success(f"Loaded **{uploaded.name}** — {len(df)} rows, {len(df.columns)} columns")

    st.markdown("### 2. Variable types")
    st.caption(
        "Stream-lite guesses numerical vs. categorical from the data. "
        "Override anything it gets wrong, and uncheck variables to exclude them."
    )

    if "var_meta" not in st.session_state or st.session_state.get("_last_file") != uploaded.name:
        var_meta = {}
        for col in df.columns:
            vtype, n, unique_n = detect_type(df[col])
            var_meta[col] = {"type": vtype, "detected": vtype, "use": True, "n": n, "unique": unique_n}
        st.session_state["var_meta"] = var_meta
        st.session_state["_last_file"] = uploaded.name

    var_meta = st.session_state["var_meta"]

    editor_df = pd.DataFrame([
        {
            "Variable": col,
            "Use": var_meta[col]["use"],
            "Detected": var_meta[col]["detected"].capitalize(),
            "Type": var_meta[col]["type"],
            "n (non-missing)": var_meta[col]["n"],
            "Unique values": var_meta[col]["unique"],
        }
        for col in df.columns
    ])

    edited = st.data_editor(
        editor_df,
        column_config={
            "Use": st.column_config.CheckboxColumn(required=True),
            "Type": st.column_config.SelectboxColumn(options=["numerical", "categorical"], required=True),
            "Detected": st.column_config.TextColumn(disabled=True),
            "Variable": st.column_config.TextColumn(disabled=True),
            "n (non-missing)": st.column_config.NumberColumn(disabled=True),
            "Unique values": st.column_config.NumberColumn(disabled=True),
        },
        hide_index=True,
        use_container_width=True,
        key="var_editor",
    )

    for _, row in edited.iterrows():
        var_meta[row["Variable"]]["use"] = bool(row["Use"])
        var_meta[row["Variable"]]["type"] = row["Type"]

    st.markdown("### 3. Grouping & test selection")
    categorical_cols = [c for c in df.columns if var_meta[c]["type"] == "categorical"]

    col1, col2, col3 = st.columns([2, 2, 1])
    with col1:
        group_col = st.selectbox(
            "Grouping variable",
            options=["— None (descriptive only) —"] + categorical_cols,
        )
        group_col = None if group_col == "— None (descriptive only) —" else group_col

        display_mode = st.radio(
            "Descriptive statistics display",
            options=["auto", "mean_sd", "median_iqr", "both"],
            format_func=lambda x: {"auto": "Auto (normality-based)",
                                    "mean_sd": "Mean \u00B1 SD",
                                    "median_iqr": "Median (IQR)",
                                    "both": "Both"}[x],
            horizontal=True,
        )

    with col2:
        test_mode = st.radio(
            "Numeric test selection",
            options=["auto", "parametric", "nonparametric"],
            format_func=lambda x: {"auto": "Auto (normality-based)",
                                    "parametric": "Force parametric",
                                    "nonparametric": "Force non-parametric"}[x],
            horizontal=True,
        )
        st.caption("Display and test selection are independent — e.g. you can show both mean\u00B1SD "
                    "and median (IQR) while still testing with Wilcoxon based on normality.")

    with col3:
        alpha = st.number_input("Significance level (\u03B1)", min_value=0.001, max_value=0.5,
                                 value=0.05, step=0.01)

    can_generate = True
    if group_col:
        levels = df[group_col].dropna().astype(str).str.strip().unique().tolist()
        counts = ", ".join(
            f"{lv} (n={(df[group_col].astype(str).str.strip() == lv).sum()})" for lv in sorted(levels)
        )
        if len(levels) < 2:
            st.warning(f"Groups in **{group_col}**: {counts} — need at least 2 groups to run comparisons.")
            can_generate = False
        else:
            st.info(f"Groups in **{group_col}**: {counts}")

    if st.button("Generate Table 1", type="primary", disabled=not can_generate):
        header, display_rows, csv_rows, flags = build_table1(df, var_meta, group_col, test_mode, alpha, display_mode)
        st.session_state["result"] = (header, display_rows, csv_rows, flags, group_col, alpha, display_mode)

    if "result" in st.session_state:
        header, display_rows, csv_rows, flags, result_group_col, result_alpha, result_display_mode = st.session_state["result"]

        st.markdown("### Table 1. Baseline characteristics")
        st.markdown(render_table_markdown(header, display_rows, result_group_col, result_alpha),
                     unsafe_allow_html=True)

        footnotes = [descriptive_footnote(result_display_mode, result_alpha)]
        if "chi2" in flags:
            footnotes.append("Chi-square test of independence used for categorical comparisons with adequate expected cell counts.")
        if "fisher" in flags:
            footnotes.append("Fisher's exact test used in place of chi-square when a 2\u00D72 table had an expected cell count below 5.")
        if "lowE" in flags:
            footnotes.append("Caution: one or more categorical comparisons above have expected cell counts below 5; chi-square approximation may be unreliable.")
        if "skipped" in flags:
            footnotes.append("Note: a statistical test could not be computed for one or more variables (e.g. insufficient data in a group) and was left blank.")
        footnotes.append(f"Bold p-values indicate statistical significance at \u03B1={result_alpha}.")
        st.caption("  \n".join(footnotes))

        dl_col1, dl_col2, dl_col3 = st.columns(3)

        with dl_col1:
            csv_buf = io.StringIO()
            pd.DataFrame(csv_rows).to_csv(csv_buf, index=False, header=False)
            st.download_button("Download CSV", csv_buf.getvalue(), file_name="table1.csv", mime="text/csv")

        with dl_col2:
            excel_buf = io.BytesIO()
            pd.DataFrame(csv_rows[1:], columns=csv_rows[0]).to_excel(excel_buf, index=False, sheet_name="Table 1")
            st.download_button("Download Excel", excel_buf.getvalue(), file_name="table1.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        with dl_col3:
            docx_bytes = build_docx(header, display_rows, result_group_col, result_alpha, flags, result_display_mode)
            st.download_button("Download Word", docx_bytes, file_name="table1.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

else:
    st.info("Upload a file to get started. Nothing leaves your machine — the app runs locally.")
